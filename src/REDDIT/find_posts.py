import pandas as pd
import os
from docx import Document
from docx.shared import Inches

# --- Configuration ---
OUTPUT_FOLDER = 'reports' 
OUTPUT_FILENAME = os.path.join(OUTPUT_FOLDER, 'sentiment_analysis_report.docx')

# --- Main Script Logic ---
try:
    df = pd.read_csv("/home/robertsory/Desktop/DESKTOP/Projects/AI_in_Education/AI_in_Education/data/processed/analyzed_sentiment.csv")
    df['selftext'] = df['selftext'].fillna('[No text body]')
except FileNotFoundError:
    print("Error: 'analyzed_sentiment.csv' not found. Please ensure the file is in the data/processed/ folder.")
    exit()

# Find the top posts
top_positive = df.sort_values(by='sentiment_score', ascending=False).head(3)
top_negative = df.sort_values(by='sentiment_score', ascending=True).head(3)
df['abs_sentiment'] = df['sentiment_score'].abs()
top_neutral = df.sort_values(by='abs_sentiment').head(3)


# --- Create and Save the Word Document ---

if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)
    print(f"Created directory: '{OUTPUT_FOLDER}'")

doc = Document()
doc.add_heading('Sentiment Analysis of Reddit Posts on AI in Education', level=0)

# --- NEW: Add the Summary Section ---
doc.add_heading('Brief Summary of Findings', level=1)
summary_text = (
    "This analysis of Reddit posts reveals a sharply polarized conversation surrounding AI in education. "
    "Negative sentiment is predominant and deeply personal, with educators expressing concerns about academic integrity, "
    "the erosion of student critical thinking, and a sense of 'moral injury' from increased workload and institutional pressures. "
    "In contrast, positive sentiment is more philosophical and forward-looking, focusing on the potential for AI to engage in logical, ethical reasoning. "
    "Neutral posts are typically factual, news-based announcements, indicating that emotional responses are tied more to the "
    "practical application of AI rather than its mere existence."
)
doc.add_paragraph(summary_text)
# --- End of New Section ---

def add_posts_to_doc(posts, category_name, doc):
    doc.add_heading(f'Top 3 Most {category_name} Posts', level=1)
    for index, row in posts.iterrows():
        doc.add_heading(row['title'], level=2)
        p = doc.add_paragraph()
        p.add_run('Subreddit: ').bold = True
        p.add_run(f"r/{row['subreddit']}\n")
        p.add_run('Sentiment Score: ').bold = True
        p.add_run(f"{row['sentiment_score']:.2f}\n")
        p.add_run('URL: ').bold = True
        p.add_run(row.get('url', 'N/A'))
        doc.add_paragraph(row['selftext'])
        doc.add_paragraph() 

add_posts_to_doc(top_positive, 'Positive 🌟', doc)
add_posts_to_doc(top_negative, 'Negative 😠', doc)
add_posts_to_doc(top_neutral, 'Neutral 😐', doc)

doc.save(OUTPUT_FILENAME)
print(f"\n✅ Successfully saved updated report to: '{OUTPUT_FILENAME}'")